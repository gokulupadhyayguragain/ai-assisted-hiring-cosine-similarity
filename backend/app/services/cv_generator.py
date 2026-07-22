"""CV Generator Service - Professional CV/Resume generation with proper formatting.

Generates professionally formatted CVs in DOCX and PDF formats with:
- Proper headings, fonts, and spacing
- Bullet points with correct indentation
- Left/center/right alignment as appropriate
- Multiple template styles (Modern, Classic, Academic, Minimal, Executive)
"""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from typing import Any


# ---------------------------------------------------------------------------
# Data Models
# ---------------------------------------------------------------------------

@dataclass
class ContactInfo:
    full_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    website: str = ""
    title: str = ""


@dataclass
class ExperienceItem:
    job_title: str = ""
    company: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    bullets: list[str] = field(default_factory=list)


@dataclass
class EducationItem:
    degree: str = ""
    institution: str = ""
    location: str = ""
    start_date: str = ""
    end_date: str = ""
    gpa: str = ""
    details: list[str] = field(default_factory=list)


@dataclass
class ProjectItem:
    name: str = ""
    description: str = ""
    technologies: str = ""
    url: str = ""
    bullets: list[str] = field(default_factory=list)


@dataclass
class CertificationItem:
    name: str = ""
    issuer: str = ""
    date: str = ""
    url: str = ""


@dataclass
class CVData:
    contact: ContactInfo = field(default_factory=ContactInfo)
    summary: str = ""
    experience: list[ExperienceItem] = field(default_factory=list)
    education: list[EducationItem] = field(default_factory=list)
    skills: dict[str, list[str]] = field(default_factory=dict)
    projects: list[ProjectItem] = field(default_factory=list)
    certifications: list[CertificationItem] = field(default_factory=list)
    languages: list[str] = field(default_factory=list)
    interests: list[str] = field(default_factory=list)
    custom_sections: dict[str, list[str]] = field(default_factory=dict)


# Template definitions
AVAILABLE_TEMPLATES = [
    {"id": "professional", "name": "Professional", "description": "Polished, versatile layout for most industries and job applications"},
    {"id": "classic", "name": "Classic", "description": "Traditional professional format, suitable for all industries"},
    {"id": "europass", "name": "Europass CV", "description": "Standard European format for EU applications and scholarships"},
    {"id": "modern", "name": "Modern", "description": "Clean design with color accents, ideal for tech and corporate roles"},
    {"id": "academic", "name": "Academic", "description": "Research-focused with emphasis on publications and education"},
    {"id": "minimal", "name": "Minimal", "description": "Simple and ATS-friendly, maximum readability"},
    {"id": "executive", "name": "Executive", "description": "Bold header with professional layout for senior roles"},
]


def _text(value: Any) -> str:
    """Return a safe string for values supplied by parsers or JSON clients."""
    return str(value).strip() if value is not None else ""


def _text_list(value: Any) -> list[str]:
    """Normalize a list field that may arrive as a string or another iterable."""
    if value is None:
        return []
    if isinstance(value, str):
        return [item.strip() for item in value.replace(";", ",").split(",") if item.strip()]
    if isinstance(value, (list, tuple, set)):
        return [_text(item) for item in value if _text(item)]
    return [_text(value)] if _text(value) else []


def _records(value: Any) -> list[dict[str, Any]]:
    """Keep only mapping records from a user/parser supplied collection."""
    return [item for item in value if isinstance(item, dict)] if isinstance(value, list) else []


def parse_cv_data(data: dict[str, Any]) -> CVData:
    """Parse raw dict (from JSON request) into CVData dataclass.

    CVs can come from hand-authored JSON, PDF extraction, or an LLM.  Those
    sources do not always agree on whether a repeated field is a list or a
    string, so normalize values here before either renderer consumes them.
    """
    contact_raw = data.get("contact") if isinstance(data.get("contact"), dict) else {}
    contact = ContactInfo(
        full_name=_text(contact_raw.get("full_name")),
        email=_text(contact_raw.get("email")),
        phone=_text(contact_raw.get("phone")),
        location=_text(contact_raw.get("location")),
        linkedin=_text(contact_raw.get("linkedin")),
        website=_text(contact_raw.get("website")),
        title=_text(contact_raw.get("title")),
    )

    experience = [
        ExperienceItem(
            job_title=_text(e.get("job_title")),
            company=_text(e.get("company")),
            location=_text(e.get("location")),
            start_date=_text(e.get("start_date")),
            end_date=_text(e.get("end_date")),
            bullets=_text_list(e.get("bullets")),
        )
        for e in _records(data.get("experience", []))
    ]

    education = [
        EducationItem(
            degree=_text(ed.get("degree")),
            institution=_text(ed.get("institution")),
            location=_text(ed.get("location")),
            start_date=_text(ed.get("start_date")),
            end_date=_text(ed.get("end_date")),
            gpa=_text(ed.get("gpa")),
            details=_text_list(ed.get("details")),
        )
        for ed in _records(data.get("education", []))
    ]

    projects = [
        ProjectItem(
            name=_text(p.get("name")),
            description=_text(p.get("description")),
            technologies=_text(p.get("technologies")),
            url=_text(p.get("url")),
            bullets=_text_list(p.get("bullets")),
        )
        for p in _records(data.get("projects", []))
    ]

    certifications = [
        CertificationItem(
            name=_text(c.get("name")),
            issuer=_text(c.get("issuer")),
            date=_text(c.get("date")),
            url=_text(c.get("url")),
        )
        for c in _records(data.get("certifications", []))
    ]

    raw_skills = data.get("skills", {})
    skills: dict[str, list[str]] = {}
    if isinstance(raw_skills, dict):
        for category, values in raw_skills.items():
            items = _text_list(values)
            if items:
                skills[_text(category) or "Skills"] = items

    raw_custom = data.get("custom_sections", {})
    custom_sections: dict[str, list[str]] = {}
    if isinstance(raw_custom, dict):
        for section, values in raw_custom.items():
            items = _text_list(values)
            if items:
                custom_sections[_text(section) or "Additional Information"] = items

    return CVData(
        contact=contact,
        summary=_text(data.get("summary")),
        experience=experience,
        education=education,
        skills=skills,
        projects=projects,
        certifications=certifications,
        languages=_text_list(data.get("languages")),
        interests=_text_list(data.get("interests")),
        custom_sections=custom_sections,
    )


# ---------------------------------------------------------------------------
# DOCX Generation
# ---------------------------------------------------------------------------

def generate_docx(cv: CVData, template: str = "modern") -> bytes:
    """Generate a professionally formatted DOCX file from CV data."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()

    # Use one predictable body font and line rhythm across every renderer.
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9.5)
    normal_style.paragraph_format.line_spacing = 1.0
    normal_style.paragraph_format.space_after = Pt(0)

    # Page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Color schemes per template
    colors = {
        "professional": {"primary": RGBColor(0, 57, 147), "secondary": RGBColor(80, 80, 80), "accent": RGBColor(0, 57, 147)},
        "modern": {"primary": RGBColor(0, 57, 147), "secondary": RGBColor(80, 80, 80), "accent": RGBColor(0, 57, 147)},
        "classic": {"primary": RGBColor(0, 0, 0), "secondary": RGBColor(60, 60, 60), "accent": RGBColor(0, 0, 0)},
        "europass": {"primary": RGBColor(0, 102, 153), "secondary": RGBColor(70, 70, 70), "accent": RGBColor(0, 102, 153)},
        "academic": {"primary": RGBColor(25, 25, 112), "secondary": RGBColor(70, 70, 70), "accent": RGBColor(25, 25, 112)},
        "minimal": {"primary": RGBColor(30, 30, 30), "secondary": RGBColor(100, 100, 100), "accent": RGBColor(30, 30, 30)},
        "executive": {"primary": RGBColor(44, 62, 80), "secondary": RGBColor(52, 73, 94), "accent": RGBColor(41, 128, 185)},
    }
    scheme = colors.get(template, colors["modern"])

    # --- HEADER / CONTACT INFO ---
    _docx_add_header(doc, cv.contact, scheme, template)

    # --- SUMMARY ---
    if cv.summary.strip():
        _docx_add_section_heading(doc, "Professional Summary", scheme)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(cv.summary.strip())
        run.font.size = Pt(9.5)
        run.font.color.rgb = scheme["secondary"]

    # --- EXPERIENCE ---
    if cv.experience:
        _docx_add_section_heading(doc, "Professional Experience", scheme)
        for exp in cv.experience:
            _docx_add_experience(doc, exp, scheme, template)

    # --- EDUCATION ---
    if cv.education:
        _docx_add_section_heading(doc, "Education", scheme)
        for edu in cv.education:
            _docx_add_education(doc, edu, scheme, template)

    # --- SKILLS ---
    if cv.skills:
        _docx_add_section_heading(doc, "Skills", scheme)
        _docx_add_skills(doc, cv.skills, scheme, template)

    # --- PROJECTS ---
    if cv.projects:
        _docx_add_section_heading(doc, "Projects", scheme)
        for proj in cv.projects:
            _docx_add_project(doc, proj, scheme, template)

    # --- CERTIFICATIONS ---
    if cv.certifications:
        _docx_add_section_heading(doc, "Certifications", scheme)
        for cert in cv.certifications:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(cert.name)
            run.bold = True
            run.font.size = Pt(10)
            if cert.issuer:
                run2 = p.add_run(f" — {cert.issuer}")
                run2.font.size = Pt(10)
                run2.font.color.rgb = scheme["secondary"]
            if cert.date:
                run3 = p.add_run(f"  ({cert.date})")
                run3.font.size = Pt(9)
                run3.font.color.rgb = scheme["secondary"]

    # --- LANGUAGES ---
    if cv.languages:
        _docx_add_section_heading(doc, "Languages", scheme)
        p = doc.add_paragraph()
        run = p.add_run("  •  ".join(cv.languages))
        run.font.size = Pt(10)
        run.font.color.rgb = scheme["secondary"]

    # --- INTERESTS ---
    if cv.interests:
        _docx_add_section_heading(doc, "Interests", scheme)
        p = doc.add_paragraph()
        run = p.add_run("  •  ".join(cv.interests))
        run.font.size = Pt(10)
        run.font.color.rgb = scheme["secondary"]

    # --- CUSTOM SECTIONS ---
    for section_name, items in cv.custom_sections.items():
        if items:
            _docx_add_section_heading(doc, section_name, scheme)
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(item)
                run.font.size = Pt(10)

    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _docx_add_header(doc, contact: ContactInfo, scheme: dict, template: str):
    """Add the name and contact header to the document."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Name
    name_para = doc.add_paragraph()
    if template in ("modern", "professional", "executive"):
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(contact.full_name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(20 if template == "executive" else 17)
    name_run.font.color.rgb = scheme["primary"]

    # Title
    if contact.title:
        title_para = doc.add_paragraph()
        title_para.alignment = name_para.alignment
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(4)
        title_run = title_para.add_run(contact.title)
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = scheme["secondary"]

    # Contact line
    contact_parts = []
    if contact.email:
        contact_parts.append(contact.email)
    if contact.phone:
        contact_parts.append(contact.phone)
    if contact.location:
        contact_parts.append(contact.location)
    if contact.linkedin:
        contact_parts.append(contact.linkedin)
    if contact.website:
        contact_parts.append(contact.website)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = name_para.alignment
        contact_para.paragraph_format.space_after = Pt(8)
        separator = "  |  " if template in ("modern", "professional", "executive") else "  •  "
        contact_run = contact_para.add_run(separator.join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = scheme["secondary"]

    # Divider line (using paragraph border for modern/professional/executive)
    if template in ("modern", "professional", "executive", "classic"):
        divider = doc.add_paragraph()
        divider.paragraph_format.space_after = Pt(6)
        # Add a thin line using underscores or em-dash
        run = divider.add_run("─" * 80)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(200, 200, 200)


def _docx_add_section_heading(doc, text: str, scheme: dict):
    """Add a formatted section heading."""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(20)
    para.paragraph_format.space_after = Pt(8)

    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = scheme["primary"]

    # Add bottom border to the paragraph
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), scheme["primary"].hex_string if hasattr(scheme["primary"], 'hex_string') else "003893")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_add_experience(doc, exp: ExperienceItem, scheme: dict, template: str):
    """Add an experience entry with title, company, dates, and bullet points."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Title + Company line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(12)

    title_run = p.add_run(exp.job_title)
    title_run.bold = True
    title_run.font.size = Pt(10.5)

    if exp.company:
        sep_run = p.add_run("  —  ")
        sep_run.font.size = Pt(11)
        sep_run.font.color.rgb = scheme["secondary"]
        comp_run = p.add_run(exp.company)
        comp_run.font.size = Pt(11)
        comp_run.font.color.rgb = scheme["accent"]

    # Date + Location line
    date_parts = []
    if exp.start_date or exp.end_date:
        date_str = f"{exp.start_date} – {exp.end_date or 'Present'}"
        date_parts.append(date_str)
    if exp.location:
        date_parts.append(exp.location)

    if date_parts:
        date_p = doc.add_paragraph()
        date_p.paragraph_format.space_after = Pt(4)
        date_p.paragraph_format.space_before = Pt(0)
        date_run = date_p.add_run("  |  ".join(date_parts))
        date_run.font.size = Pt(9)
        date_run.font.italic = True
        date_run.font.color.rgb = scheme["secondary"]

    # Bullet points
    for bullet in exp.bullets:
        if bullet.strip():
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.28)
            bp.paragraph_format.first_line_indent = Inches(-0.16)
            bp.paragraph_format.right_indent = Inches(0)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.line_spacing = 1.0
            bp.paragraph_format.keep_together = True
            bp.paragraph_format.widow_control = True
            run = bp.add_run(bullet.strip())
            run.font.name = "Arial"
            run.font.size = Pt(9.5)


def _docx_add_education(doc, edu: EducationItem, scheme: dict, template: str):
    """Add an education entry."""
    from docx.shared import Pt, Inches

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

    deg_run = p.add_run(edu.degree)
    deg_run.bold = True
    deg_run.font.size = Pt(11)

    if edu.institution:
        sep_run = p.add_run("  —  ")
        sep_run.font.size = Pt(11)
        inst_run = p.add_run(edu.institution)
        inst_run.font.size = Pt(11)
        inst_run.font.color.rgb = scheme["accent"]

    # Date and location
    meta_parts = []
    if edu.start_date or edu.end_date:
        meta_parts.append(f"{edu.start_date} – {edu.end_date or 'Present'}")
    if edu.location:
        meta_parts.append(edu.location)
    if edu.gpa:
        meta_parts.append(f"GPA: {edu.gpa}")

    if meta_parts:
        mp = doc.add_paragraph()
        mp.paragraph_format.space_after = Pt(3)
        mp.paragraph_format.space_before = Pt(0)
        run = mp.add_run("  |  ".join(meta_parts))
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = scheme["secondary"]

    # Details
    for detail in edu.details:
        if detail.strip():
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.28)
            bp.paragraph_format.first_line_indent = Inches(-0.16)
            bp.paragraph_format.right_indent = Inches(0)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.line_spacing = 1.0
            bp.paragraph_format.keep_together = True
            bp.paragraph_format.widow_control = True
            run = bp.add_run(detail.strip())
            run.font.name = "Arial"
            run.font.size = Pt(9.5)


def _docx_add_skills(doc, skills: dict[str, list[str]], scheme: dict, template: str):
    """Add skills section with category grouping."""
    from docx.shared import Pt

    if isinstance(skills, dict):
        for category, skill_list in skills.items():
            if skill_list:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(5)
                cat_run = p.add_run(f"{category}: ")
                cat_run.bold = True
                cat_run.font.size = Pt(10)
                skills_run = p.add_run(", ".join(skill_list))
                skills_run.font.size = Pt(10)
                skills_run.font.color.rgb = scheme["secondary"]
    elif isinstance(skills, list):
        p = doc.add_paragraph()
        run = p.add_run(", ".join(skills))
        run.font.size = Pt(10)


def _docx_add_project(doc, proj: ProjectItem, scheme: dict, template: str):
    """Add a project entry."""
    from docx.shared import Pt, Inches

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)

    name_run = p.add_run(proj.name)
    name_run.bold = True
    name_run.font.size = Pt(10.5)

    if proj.technologies:
        tech_run = p.add_run(f"  [{proj.technologies}]")
        tech_run.font.size = Pt(9)
        tech_run.font.italic = True
        tech_run.font.color.rgb = scheme["secondary"]

    if proj.description:
        desc_p = doc.add_paragraph()
        desc_p.paragraph_format.space_after = Pt(3)
        desc_run = desc_p.add_run(proj.description)
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = scheme["secondary"]

    for bullet in proj.bullets:
        if bullet.strip():
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.28)
            bp.paragraph_format.first_line_indent = Inches(-0.16)
            bp.paragraph_format.right_indent = Inches(0)
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(2)
            bp.paragraph_format.line_spacing = 1.0
            bp.paragraph_format.keep_together = True
            bp.paragraph_format.widow_control = True
            run = bp.add_run(bullet.strip())
            run.font.name = "Arial"
            run.font.size = Pt(9.5)


# ---------------------------------------------------------------------------
# PDF Generation (using reportlab)
# ---------------------------------------------------------------------------

def generate_pdf(cv: CVData, template: str = "modern") -> bytes:
    """Generate a professionally formatted PDF file from CV data."""
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable, ListFlowable, ListItem,
    )
    from xml.sax.saxutils import escape

    # Color schemes
    color_schemes = {
        "professional": {"primary": colors.HexColor("#003893"), "secondary": colors.HexColor("#505050"), "accent": colors.HexColor("#003893"), "line": colors.HexColor("#003893")},
        "modern": {"primary": colors.HexColor("#003893"), "secondary": colors.HexColor("#505050"), "accent": colors.HexColor("#003893"), "line": colors.HexColor("#003893")},
        "classic": {"primary": colors.HexColor("#000000"), "secondary": colors.HexColor("#3c3c3c"), "accent": colors.HexColor("#000000"), "line": colors.HexColor("#333333")},
        "europass": {"primary": colors.HexColor("#006699"), "secondary": colors.HexColor("#464646"), "accent": colors.HexColor("#006699"), "line": colors.HexColor("#006699")},
        "academic": {"primary": colors.HexColor("#191970"), "secondary": colors.HexColor("#464646"), "accent": colors.HexColor("#191970"), "line": colors.HexColor("#191970")},
        "minimal": {"primary": colors.HexColor("#1e1e1e"), "secondary": colors.HexColor("#646464"), "accent": colors.HexColor("#1e1e1e"), "line": colors.HexColor("#cccccc")},
        "executive": {"primary": colors.HexColor("#2c3e50"), "secondary": colors.HexColor("#34495e"), "accent": colors.HexColor("#2980b9"), "line": colors.HexColor("#2c3e50")},
    }
    scheme = color_schemes.get(template, color_schemes["modern"])

    # Define styles
    styles = getSampleStyleSheet()

    name_align = TA_LEFT if template in ("modern", "professional", "executive") else TA_CENTER

    style_name = ParagraphStyle(
        "CVName", parent=styles["Title"],
        fontSize=20 if template == "executive" else 16,
        alignment=name_align,
        textColor=scheme["primary"],
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )
    style_title = ParagraphStyle(
        "CVTitle", parent=styles["Normal"],
        fontSize=11,
        alignment=name_align,
        textColor=scheme["secondary"],
        spaceAfter=4,
        fontName="Helvetica",
    )
    style_contact = ParagraphStyle(
        "CVContact", parent=styles["Normal"],
        fontSize=9,
        alignment=name_align,
        textColor=scheme["secondary"],
        spaceAfter=12,
        fontName="Helvetica",
    )
    style_section = ParagraphStyle(
        "CVSection", parent=styles["Heading2"],
        fontSize=10.5,
        leading=12,
        textColor=scheme["primary"],
        spaceBefore=12,
        spaceAfter=4,
        fontName="Helvetica-Bold",
        keepWithNext=True,
    )
    style_subsection = ParagraphStyle(
        "CVSubsection", parent=styles["Normal"],
        fontSize=10,
        leading=11.5,
        textColor=colors.black,
        spaceBefore=7,
        spaceAfter=1,
        fontName="Helvetica-Bold",
        keepWithNext=True,
    )
    style_meta = ParagraphStyle(
        "CVMeta", parent=styles["Normal"],
        fontSize=8.5,
        leading=10,
        textColor=scheme["secondary"],
        spaceAfter=2,
        fontName="Helvetica-Oblique",
    )
    style_body = ParagraphStyle(
        "CVBody", parent=styles["Normal"],
        fontSize=9.5,
        textColor=scheme["secondary"],
        spaceAfter=4,
        leading=11.5,
        fontName="Helvetica",
        wordWrap="LTR",
        splitLongWords=True,
    )
    style_bullet = ParagraphStyle(
        "CVBullet", parent=styles["Normal"],
        fontSize=9.2,
        textColor=colors.HexColor("#333333"),
        leftIndent=14,
        firstLineIndent=-10,
        rightIndent=0,
        spaceBefore=0,
        spaceAfter=2,
        leading=11.2,
        fontName="Helvetica",
        wordWrap="LTR",
        splitLongWords=True,
        allowWidows=0,
        allowOrphans=0,
    )
    style_skill_cat = ParagraphStyle(
        "CVSkillCat", parent=styles["Normal"],
        fontSize=9.5,
        textColor=colors.black,
        spaceAfter=3,
        leading=11.5,
        fontName="Helvetica",
        wordWrap="LTR",
        splitLongWords=True,
    )

    # Build story
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=1.5 * cm, leftMargin=1.5 * cm,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
        title=f"{cv.contact.full_name or 'CV'} - Resume",
    )

    story: list = []

    # --- HEADER ---
    story.append(Paragraph(escape(cv.contact.full_name or "Your Name"), style_name))

    if cv.contact.title:
        story.append(Paragraph(escape(cv.contact.title), style_title))

    contact_parts = []
    if cv.contact.email:
        contact_parts.append(cv.contact.email)
    if cv.contact.phone:
        contact_parts.append(cv.contact.phone)
    if cv.contact.location:
        contact_parts.append(cv.contact.location)
    if cv.contact.linkedin:
        contact_parts.append(cv.contact.linkedin)
    if cv.contact.website:
        contact_parts.append(cv.contact.website)

    if contact_parts:
        separator = "  |  " if template in ("modern", "professional", "executive") else "  •  "
        story.append(Paragraph(escape(separator.join(contact_parts)), style_contact))

    # Divider
    story.append(HRFlowable(
        width="100%", thickness=1.2 if template in ("modern", "professional", "executive") else 0.5,
        color=scheme["line"], spaceAfter=12,
    ))

    # --- SUMMARY ---
    if cv.summary.strip():
        story.append(Paragraph("PROFESSIONAL SUMMARY", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))
        story.append(Paragraph(escape(cv.summary.strip()), style_body))

    # --- EXPERIENCE ---
    if cv.experience:
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))

        for exp in cv.experience:
            # Title + Company
            title_text = f"<b>{escape(exp.job_title)}</b>"
            if exp.company:
                title_text += f'  —  <font color="{scheme["accent"].hexval()}">{escape(exp.company)}</font>'
            story.append(Paragraph(title_text, style_subsection))

            # Date + Location
            meta_parts = []
            if exp.start_date or exp.end_date:
                meta_parts.append(f"{exp.start_date} – {exp.end_date or 'Present'}")
            if exp.location:
                meta_parts.append(exp.location)
            if meta_parts:
                story.append(Paragraph(escape("  |  ".join(meta_parts)), style_meta))

            # Bullets
            for bullet in exp.bullets:
                if bullet.strip():
                    story.append(Paragraph(f"•  {escape(bullet.strip())}", style_bullet))

    # --- EDUCATION ---
    if cv.education:
        story.append(Paragraph("EDUCATION", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))

        for edu in cv.education:
            title_text = f"<b>{escape(edu.degree)}</b>"
            if edu.institution:
                title_text += f'  —  <font color="{scheme["accent"].hexval()}">{escape(edu.institution)}</font>'
            story.append(Paragraph(title_text, style_subsection))

            meta_parts = []
            if edu.start_date or edu.end_date:
                meta_parts.append(f"{edu.start_date} – {edu.end_date or 'Present'}")
            if edu.location:
                meta_parts.append(edu.location)
            if edu.gpa:
                meta_parts.append(f"GPA: {edu.gpa}")
            if meta_parts:
                story.append(Paragraph(escape("  |  ".join(meta_parts)), style_meta))

            for detail in edu.details:
                if detail.strip():
                    story.append(Paragraph(f"•  {escape(detail.strip())}", style_bullet))

    # --- SKILLS ---
    if cv.skills:
        story.append(Paragraph("SKILLS", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))

        if isinstance(cv.skills, dict):
            for category, skill_list in cv.skills.items():
                if skill_list:
                    text = f"<b>{escape(category)}:</b>  {escape(', '.join(skill_list))}"
                    story.append(Paragraph(text, style_skill_cat))
        elif isinstance(cv.skills, list):
            story.append(Paragraph(escape(", ".join(cv.skills)), style_body))

    # --- PROJECTS ---
    if cv.projects:
        story.append(Paragraph("PROJECTS", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))

        for proj in cv.projects:
            title_text = f"<b>{escape(proj.name)}</b>"
            if proj.technologies:
                title_text += f"  <i>[{escape(proj.technologies)}]</i>"
            story.append(Paragraph(title_text, style_subsection))

            if proj.description:
                story.append(Paragraph(escape(proj.description), style_meta))

            for bullet in proj.bullets:
                if bullet.strip():
                    story.append(Paragraph(f"•  {escape(bullet.strip())}", style_bullet))

    # --- CERTIFICATIONS ---
    if cv.certifications:
        story.append(Paragraph("CERTIFICATIONS", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))

        for cert in cv.certifications:
            text = f"<b>{escape(cert.name)}</b>"
            if cert.issuer:
                text += f" — {escape(cert.issuer)}"
            if cert.date:
                text += f"  ({escape(cert.date)})"
            story.append(Paragraph(text, style_skill_cat))

    # --- LANGUAGES ---
    if cv.languages:
        story.append(Paragraph("LANGUAGES", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))
        story.append(Paragraph(escape("  •  ".join(cv.languages)), style_body))

    # --- INTERESTS ---
    if cv.interests:
        story.append(Paragraph("INTERESTS", style_section))
        story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))
        story.append(Paragraph(escape("  •  ".join(cv.interests)), style_body))

    # --- CUSTOM SECTIONS ---
    for section_name, items in cv.custom_sections.items():
        if items:
            story.append(Paragraph(escape(section_name.upper()), style_section))
            story.append(HRFlowable(width="100%", thickness=0.3, color=scheme["line"], spaceAfter=8))
            for item in items:
                story.append(Paragraph(f"•  {escape(item)}", style_bullet))

    doc.build(story)
    return buffer.getvalue()
