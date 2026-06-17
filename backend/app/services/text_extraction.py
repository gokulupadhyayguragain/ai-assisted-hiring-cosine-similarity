from __future__ import annotations

from io import BytesIO
from pathlib import Path

from fastapi import UploadFile

from backend.app.models import ExtractedDocument


async def extract_upload(file: UploadFile) -> ExtractedDocument:
    raw = await file.read()
    filename = file.filename or "uploaded-file"
    suffix = Path(filename).suffix.lower()
    warnings: list[str] = []

    if suffix in {".txt", ".md", ".csv"}:
        return ExtractedDocument(filename=filename, text=_decode_bytes(raw), file_type=suffix.lstrip("."))

    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}:
        text, image_warnings = _extract_image(raw)
        warnings.extend(image_warnings)
        return ExtractedDocument(filename=filename, text=text, file_type=suffix.lstrip("."), warnings=warnings)

    if suffix == ".pdf":
        text, pdf_warnings = _extract_pdf(raw)
        warnings.extend(pdf_warnings)
        return ExtractedDocument(filename=filename, text=text, file_type="pdf", warnings=warnings)

    if suffix == ".docx":
        text, docx_warnings = _extract_docx(raw)
        warnings.extend(docx_warnings)
        return ExtractedDocument(filename=filename, text=text, file_type="docx", warnings=warnings)

    warnings.append(f"Unsupported extension {suffix or '(none)'}; attempted plain-text decoding.")
    return ExtractedDocument(filename=filename, text=_decode_bytes(raw), file_type=suffix.lstrip(".") or "unknown", warnings=warnings)


def _decode_bytes(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _extract_pdf(raw: bytes) -> tuple[str, list[str]]:
    try:
        from pypdf import PdfReader
    except Exception:
        return _extract_pdf_with_ocr(raw, ["pypdf is not installed; using OCR fallback for PDF."])

    try:
        reader = PdfReader(BytesIO(raw))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(page.strip() for page in pages if page.strip())
        if not text:
            return _extract_pdf_with_ocr(raw, ["No selectable text found in PDF; using OCR fallback."])
        return text, []
    except Exception as exc:
        return _extract_pdf_with_ocr(raw, [f"PDF extraction failed ({exc}); using OCR fallback."])


def _extract_pdf_with_ocr(raw: bytes, warnings: list[str]) -> tuple[str, list[str]]:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:
        warnings.append(f"PyMuPDF is not installed ({exc}); PDF was decoded as plain text.")
        return _decode_bytes(raw), warnings

    text_parts: list[str] = []
    try:
        document = fitz.open(stream=raw, filetype="pdf")
    except Exception as exc:
        warnings.append(f"PDF open failed for OCR ({exc}); decoded as plain text.")
        return _decode_bytes(raw), warnings

    for index, page in enumerate(document, start=1):
        try:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = _pixmap_to_image(pixmap)
            page_text = _ocr_image(image)
            if page_text.strip():
                text_parts.append(page_text.strip())
        except Exception as exc:
            warnings.append(f"OCR failed on page {index} ({exc}).")

    text = "\n".join(part for part in text_parts if part)
    if text:
        warnings.append("OCR used for PDF pages without selectable text.")
        return text, warnings

    warnings.append("No extractable text found after OCR; decoded as plain text.")
    return _decode_bytes(raw), warnings


def _extract_image(raw: bytes) -> tuple[str, list[str]]:
    try:
        from PIL import Image
    except Exception as exc:
        return _decode_bytes(raw), [f"Pillow is not installed ({exc}); image was decoded as plain text."]

    try:
        image = Image.open(BytesIO(raw))
        text = _ocr_image(image)
        if text.strip():
            return text, ["OCR used for image document."]
        return _decode_bytes(raw), ["No text detected in image; decoded as plain text."]
    except Exception as exc:
        return _decode_bytes(raw), [f"Image OCR failed ({exc}); decoded as plain text."]


def _ocr_image(image) -> str:
    try:
        import pytesseract
    except Exception as exc:
        raise RuntimeError(f"pytesseract is not installed ({exc})") from exc

    try:
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as exc:
        raise RuntimeError(f"tesseract OCR failed ({exc})") from exc


def _pixmap_to_image(pixmap):
    try:
        from PIL import Image
    except Exception as exc:
        raise RuntimeError(f"Pillow is not installed ({exc})") from exc

    mode = "RGB" if pixmap.alpha == 0 else "RGBA"
    return Image.frombytes(mode, [pixmap.width, pixmap.height], pixmap.samples)


def _extract_docx(raw: bytes) -> tuple[str, list[str]]:
    try:
        from docx import Document
    except Exception:
        return _decode_bytes(raw), ["python-docx is not installed; DOCX was decoded as plain text."]

    try:
        document = Document(BytesIO(raw))
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts), []
    except Exception as exc:
        return _decode_bytes(raw), [f"DOCX extraction failed ({exc}); decoded as plain text."]
