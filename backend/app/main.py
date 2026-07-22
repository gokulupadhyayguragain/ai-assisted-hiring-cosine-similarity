from __future__ import annotations

from pathlib import Path
from typing import Annotated
from datetime import UTC, datetime
import os
import re
import uuid
from urllib.parse import urlencode

# Load environment variables from the repo-root .env for local (non-docker) runs.
# In docker, compose injects these via env_file/environment; load_dotenv does not
# override already-set variables, so docker values always take precedence.
from dotenv import load_dotenv
load_dotenv(Path(__file__).resolve().parents[2] / ".env")

from fastapi import FastAPI, File, Form, HTTPException, UploadFile, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import PlainTextResponse, Response, RedirectResponse
from pydantic import BaseModel, Field

from backend.app.models import ExtractedDocument, to_plain
from backend.app.services.matching import MatchingEngine
from backend.app.services.embeddings import precompute_embeddings
from backend.app.services.reporting import candidate_report_pdf, session_to_csv, resume_to_pdf
from backend.app.services.storage import SessionStore
from backend.app.services.text_extraction import extract_upload
from backend.app.services.comparison import compare_resumes
from backend.app.services.auth_handler import (
    SELF_SERVICE_ROLES,
    create_user,
    authenticate,
    get_user_from_token,
    get_user_by_email,
    update_user,
    issue_token,
    generate_verification,
    verify_email as verify_email_record,
)
from backend.app.services.job_board import JobBoardStore, ExternalJob
from backend.app.services.adzuna_client import import_recent_jobs as adzuna_import
from backend.app.services.greenhouse_client import import_company_jobs as greenhouse_import
from backend.app.services.sso_handler import (
    APP_URL,
    google_configured,
    build_google_auth_url,
    google_oauth_callback,
    send_verification_email,
)
from backend.app.services.cv_optimizer import (
    optimize_cv, get_template_names, get_scholarship_types,
    build_optimized_cv_data, calculate_skill_gap,
)
from backend.app.services.career_guide import generate_action_plan
from backend.app.services.cv_generator import (
    generate_pdf, generate_docx, parse_cv_data, AVAILABLE_TEMPLATES, CVData,
)
from backend.app.services import llm_service
import logging
import threading
import zipfile
import tarfile



logger = logging.getLogger(__name__)


APP_ROOT = Path(__file__).resolve().parents[1]
STORE = SessionStore(APP_ROOT / "runtime" / "sessions")
MODEL_DIR = APP_ROOT / "runtime" / "models"
JOB_BOARD = JobBoardStore(APP_ROOT / "runtime")

# Lazy-initialized matching engine — loads on first use so uvicorn starts
# immediately instead of blocking on HuggingFace model download.
_ENGINE_INSTANCE = None
_ENGINE_LOCK = threading.Lock()

def get_engine() -> MatchingEngine:
    global _ENGINE_INSTANCE
    if _ENGINE_INSTANCE is None:
        with _ENGINE_LOCK:
            if _ENGINE_INSTANCE is None:
                _ENGINE_INSTANCE = MatchingEngine()
    return _ENGINE_INSTANCE

app = FastAPI(
    title="AI Assisted Hiring using Cosine Similarity",
    description=(
        "FastAPI backend for resume/JD extraction, anonymization, hybrid TF-IDF plus semantic "
        "cosine scoring, explainable candidate ranking, exports, and bias auditing."
    ),
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Auth dependency
# ---------------------------------------------------------------------------

async def get_current_user(authorization: str | None = Header(None)) -> dict:
    """Extract and validate the current user from the Authorization header."""
    if not authorization:
        raise HTTPException(status_code=401, detail="Authorization header required")
    scheme, _, token = authorization.partition(" ")
    if scheme.lower() != "bearer" or not token:
        raise HTTPException(status_code=401, detail="Bearer token required")
    user = get_user_from_token(token)
    if user is None:
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    return user



async def get_optional_user(authorization: str | None = Header(None)) -> dict | None:
    """Return the signed-in user when present, without blocking public analysis."""
    if not authorization:
        return None
    scheme, _, token = authorization.partition(" ")
    if scheme.lower() != "bearer" or not token:
        return None
    return get_user_from_token(token)

class JobCreate(BaseModel):
    title: str = Field(min_length=1, max_length=180)
    department: str = Field(default="", max_length=120)
    experience: str = Field(default="", max_length=80)
    location: str = Field(default="", max_length=120)
    salary: str = Field(default="", max_length=120)
    required_skills: list[str] = Field(default_factory=list)
    description: str = Field(default="", max_length=20000)
    created_by: str = Field(default="recruiter", max_length=40)


class JobRecord(JobCreate):
    job_id: str
    created_at: str
    updated_at: str


class EmbeddingPrecomputeRequest(BaseModel):
    texts: list[str] = Field(min_length=1)


class EmbeddingPrecomputeResponse(BaseModel):
    cached: list[dict[str, object]]


@app.get("/")
def root() -> dict[str, str]:
    return {
        "name": "AI Assisted Hiring using Cosine Similarity",
        "docs": "/docs",
        "health": "/health",
    }


@app.get("/health")
def health() -> dict[str, object]:
    # Report the engine status without triggering a lazy load. The engine is
    # created on the first call that actually needs it (e.g. /api/analyze),
    # which may download the BGE model from HuggingFace. Until then the
    # server stays responsive even if HuggingFace is slow or unreachable.
    if _ENGINE_INSTANCE is not None:
        eng = _ENGINE_INSTANCE
        return {
            "status": "ok",
            "semantic_enabled": eng.semantic_enabled,
            "semantic_model": eng.semantic_model_name if eng.semantic_enabled else "not loaded yet",
        }
    return {
        "status": "ok",
        "semantic_enabled": False,
        "semantic_model": "not loaded yet — will load on first analysis",
    }


@app.post("/api/jobs", response_model=JobRecord)
def create_job(payload: JobCreate) -> dict:
    now = datetime.now(UTC).isoformat()
    job = {
        "job_id": uuid.uuid4().hex[:12],
        "created_at": now,
        "updated_at": now,
        "title": payload.title.strip(),
        "department": payload.department.strip(),
        "experience": payload.experience.strip(),
        "location": payload.location.strip(),
        "salary": payload.salary.strip(),
        "required_skills": [item.strip() for item in payload.required_skills if item.strip()],
        "description": payload.description.strip(),
        "created_by": payload.created_by.strip() or "recruiter",
    }
    STORE.save_job(job)
    return job


@app.get("/api/jobs", response_model=list[JobRecord])
def list_jobs(limit: int = 20) -> list[dict]:
    bounded = min(max(limit, 1), 100)
    return STORE.list_jobs(limit=bounded)


@app.get("/api/jobs/{job_id}", response_model=JobRecord)
def get_job(job_id: str) -> dict:
    job = STORE.load_job(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    return job


@app.put("/api/jobs/{job_id}", response_model=JobRecord)
def update_job(job_id: str, payload: JobCreate) -> dict:
    existing = STORE.load_job(job_id)
    if existing is None:
        raise HTTPException(status_code=404, detail="Job not found.")

    now = datetime.now(UTC).isoformat()
    job = {
        "job_id": job_id,
        "created_at": existing.get("created_at", now),
        "updated_at": now,
        "title": payload.title.strip(),
        "department": payload.department.strip(),
        "experience": payload.experience.strip(),
        "location": payload.location.strip(),
        "salary": payload.salary.strip(),
        "required_skills": [item.strip() for item in payload.required_skills if item.strip()],
        "description": payload.description.strip(),
        "created_by": payload.created_by.strip() or existing.get("created_by", "recruiter"),
    }
    STORE.save_job(job)
    return job


# ---------------------------------------------------------------------------
# Candidate resume library and skill-gap analysis
# ---------------------------------------------------------------------------

class SkillGapRequest(BaseModel):
    cv_text: str = Field(min_length=10)
    target_text: str = Field(min_length=10)
    resume_id: str = Field(default="", max_length=40)
    target_label: str = Field(default="Job description", max_length=180)


@app.post("/api/candidate/resumes")
async def upload_candidate_resume(
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user),
) -> dict:
    """Store a candidate resume for later analysis and report history."""
    document = await extract_upload(file)
    if not document.text.strip():
        raise HTTPException(status_code=400, detail="No extractable text found in the uploaded resume.")
    record = STORE.save_resume(
        current_user["user_id"],
        document.filename,
        document.text,
        document.file_type,
    )
    record.pop("text", None)
    return record


@app.get("/api/candidate/resumes")
def list_candidate_resumes(
    limit: int = 50,
    current_user: dict = Depends(get_current_user),
) -> dict:
    return {"resumes": STORE.list_resumes(current_user["user_id"], limit=limit)}


@app.get("/api/candidate/resumes/{resume_id}")
def get_candidate_resume(
    resume_id: str,
    current_user: dict = Depends(get_current_user),
) -> dict:
    record = STORE.load_resume(current_user["user_id"], resume_id)
    if record is None:
        raise HTTPException(status_code=404, detail="Resume not found.")
    return record


@app.delete("/api/candidate/resumes/{resume_id}")
def delete_candidate_resume(
    resume_id: str,
    current_user: dict = Depends(get_current_user),
) -> dict:
    if not STORE.delete_resume(current_user["user_id"], resume_id):
        raise HTTPException(status_code=404, detail="Resume not found.")
    return {"deleted": True, "resume_id": resume_id}


@app.post("/api/cv/skill-gap")
def candidate_skill_gap(
    payload: SkillGapRequest,
    current_user: dict | None = Depends(get_optional_user),
) -> dict:
    """Return explicit required-skill coverage and persist it when signed in."""
    report = calculate_skill_gap(payload.cv_text, payload.target_text)
    report.update({
        "analysis_type": "skill-gap",
        "target_label": payload.target_label,
        "created_at": datetime.now(UTC).isoformat(),
    })
    if current_user and payload.resume_id:
        saved = STORE.save_resume_analysis(current_user["user_id"], payload.resume_id, report)
        if saved is None:
            raise HTTPException(status_code=404, detail="Resume not found.")
    return report


@app.delete("/api/jobs/{job_id}")
def delete_job(job_id: str) -> dict:
    existing = STORE.load_job(job_id)
    if existing is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    STORE.delete_job(job_id)
    return {"deleted": True, "job_id": job_id}


# ---------------------------------------------------------------------------
# CV Optimization API Routes
# ---------------------------------------------------------------------------


class CvOptimizeRequest(BaseModel):
    cv_text: str = Field(min_length=10)
    target_text: str = Field(min_length=10)
    mode: str = Field(default="job", pattern="^(job|scholarship)$")
    template: str = Field(default="modern", max_length=30)
    name: str = Field(default="", max_length=100)
    email: str = Field(default="", max_length=120)
    phone: str = Field(default="", max_length=30)
    location: str = Field(default="", max_length=120)
    title: str = Field(default="", max_length=100)
    institution: str = Field(default="", max_length=120)
    scholarship_type: str = Field(default="", max_length=30)
    suggestions: list[str] = Field(default_factory=list, max_length=20)
    resume_id: str = Field(default="", max_length=40)


@app.post("/api/cv/optimize")
def cv_optimize(
    payload: CvOptimizeRequest,
    current_user: dict | None = Depends(get_optional_user),
) -> dict:
    """Optimize a CV against a job description or scholarship criteria.

    Returns ATS score, keyword analysis, suggestions, before/after comparison,
    and a formatted CV using the selected template.
    """
    personal_info = {
        "name": payload.name or "Full Name",
        "email": payload.email or "email@example.com",
        "phone": payload.phone or "",
        "location": payload.location or "",
        "title": payload.title or "",
        "institution": payload.institution or "",
    }

    result = optimize_cv(
        cv_text=payload.cv_text,
        target_text=payload.target_text,
        mode=payload.mode,
        template=payload.template,
        personal_info=personal_info,
        scholarship_type=payload.scholarship_type,
        additional_suggestions=payload.suggestions,
        use_generation_model=True,
    )

    import dataclasses
    plain = dataclasses.asdict(result)
    if current_user and payload.resume_id:
        STORE.save_resume_analysis(current_user["user_id"], payload.resume_id, {
            "analysis_type": "optimization",
            "target_label": "CV optimization",
            "created_at": plain["created_at"],
            "ats_score": plain["ats_score"],
            "skill_coverage": plain["skill_coverage"],
            "skill_gap_percentage": plain["skill_gap_percentage"],
            "matched_skill_count": len(plain["matched_skills"]),
            "required_skill_count": len(plain["required_skills"]),
            "generation_source": plain["generation_source"],
        })
    return plain


@app.get("/api/cv/templates")
def list_cv_templates() -> dict:
    """List available CV templates."""
    return {"templates": get_template_names()}


@app.get("/api/cv/scholarship-types")
def list_scholarship_types() -> dict:
    """List available scholarship types for CV optimization."""
    return {"types": get_scholarship_types()}


@app.get("/api/llm/status")
def llm_status() -> dict:
    """Check if the LLM model is available and loaded."""
    return {
        "available": llm_service.is_available(),
        "model": "Qwen2.5-0.5B-Instruct (Q4_K_M)",
        "ram_usage": "~400MB",
    }


class LlmRestructureRequest(BaseModel):
    cv_text: str = Field(..., min_length=50)
    job_description: str = Field(default="")


@app.post("/api/cv/restructure")
def cv_restructure_with_llm(payload: LlmRestructureRequest) -> dict:
    """Use LLM to intelligently restructure a CV from raw text.

    This endpoint uses a lightweight on-device LLM to understand and
    properly structure CV content, even from poorly-formatted PDF extractions.
    Falls back to rule-based parsing if LLM is unavailable.
    """
    # Try LLM
    result = llm_service.restructure_cv(payload.cv_text, payload.job_description)
    if result:
        return {"source": "llm", "data": result}

    # Fallback to rule-based
    from backend.app.services.cv_parser import parse_cv_text
    result = parse_cv_text(payload.cv_text)
    return {"source": "rule-based", "data": result}


class ActionPlanRequest(BaseModel):
    """Request body for generating a career action plan."""
    matched_skills: list[str] = Field(default_factory=list)
    missing_skills: list[str] = Field(default_factory=list)
    missing_keywords: list[str] = Field(default_factory=list)
    mode: str = Field(default="job", max_length=20)
    target_role: str = Field(default="", max_length=100)


@app.post("/api/cv/action-plan")
def cv_action_plan(payload: ActionPlanRequest) -> dict:
    """Generate a personalized career development action plan.

    Takes skill gap analysis results and returns learning paths with curated
    resources, project suggestions with step-by-step guides, application tips,
    and a recommended timeline to become job-ready.
    """
    result = generate_action_plan(
        matched_skills=payload.matched_skills,
        missing_skills=payload.missing_skills,
        missing_keywords=payload.missing_keywords,
        mode=payload.mode,
        target_role=payload.target_role,
    )

    import dataclasses
    return dataclasses.asdict(result)


class OptimizedCvDownloadRequest(BaseModel):
    """Request to generate a formatted PDF/DOCX from optimized CV content."""
    cv_text: str = Field(..., min_length=10)
    target_text: str = Field(default="")
    mode: str = Field(default="job")
    template: str = Field(default="modern", max_length=30)
    format: str = Field(default="pdf", pattern="^(pdf|docx)$")
    name: str = Field(default="", max_length=100)
    email: str = Field(default="", max_length=120)
    phone: str = Field(default="", max_length=30)
    location: str = Field(default="", max_length=120)
    title: str = Field(default="", max_length=100)
    current_skills: list[str] = Field(default_factory=list)
    suggestions: list[str] = Field(default_factory=list, max_length=20)
    apply_suggestions: bool = False
    missing_info: dict[str, str] = Field(default_factory=dict, max_length=30)


@app.post("/api/cv/optimize/download")
def cv_optimize_download(payload: OptimizedCvDownloadRequest):
    """Optimize a CV against a JD and return a real formatted PDF or DOCX file.

    Uses LLM for intelligent parsing when available, falls back to rule-based.
    1. Parses the original CV text into structured sections
    2. Runs optimization to identify keywords/skills to inject
    3. Generates a professional document using the selected template
    """
    from backend.app.services.cv_parser import parse_cv_text
    from backend.app.services.cv_optimizer import optimize_cv

    # Step 1: Parse with the lossless rule-based parser first. The optional
    # LLM can improve a completely unstructured document, but it must not
    # replace a parser result because doing so can drop projects, certificates,
    # or custom sections from an otherwise valid CV.
    cv_payload = parse_cv_text(payload.cv_text)
    structured_count = sum(
        bool(cv_payload.get(key))
        for key in ("summary", "experience", "education", "skills", "projects", "certifications")
    )
    if structured_count == 0:
        try:
            llm_result = llm_service.restructure_cv(payload.cv_text, payload.target_text)
            if llm_result:
                cv_payload = llm_result
                logger.info("CV parsed via LLM fallback")
        except Exception as e:
            logger.warning(f"LLM parsing fallback failed; using raw parser result: {e}")
    else:
        logger.info("CV parsed via lossless rule-based parser")

    contact = dict(cv_payload.get("contact") or {})
    cv_payload["contact"] = contact

    # Apply optional year-only answers supplied by the candidate. Values are
    # deliberately limited to a year or Present so the optimizer never asks
    # for unnecessary month-level precision.
    for key, value in payload.missing_info.items():
        clean_value = str(value).strip()
        if not re.fullmatch(r"(?:\d{4}|Present|Current)", clean_value, re.IGNORECASE):
            continue
        match = re.fullmatch(r"(experience|education)_(\d+)_(start_date|end_date)", key)
        if not match:
            continue
        collection = cv_payload.get(match.group(1))
        index = int(match.group(2))
        if isinstance(collection, list) and index < len(collection) and isinstance(collection[index], dict):
            collection[index][match.group(3)] = clean_value

    # Override contact with request fields if provided
    if payload.name:
        contact["full_name"] = payload.name
    if payload.email:
        contact["email"] = payload.email
    if payload.phone:
        contact["phone"] = payload.phone
    if payload.location:
        contact["location"] = payload.location
    if payload.title:
        contact["title"] = payload.title

    # Step 2: Run optimization to get skills to inject
    personal_info = {
        "name": contact.get("full_name", "Candidate"),
        "email": contact.get("email", ""),
        "phone": contact.get("phone", ""),
        "location": contact.get("location", ""),
        "title": contact.get("title", ""),
        "institution": "",
    }

    opt_result = optimize_cv(
        cv_text=payload.cv_text,
        target_text=payload.target_text,
        mode=payload.mode,
        template=payload.template,
        personal_info=personal_info,
        additional_suggestions=payload.suggestions,
    )

    # Step 3: Deterministic generation keeps facts, improves wording and
    # prioritizes evidence. Missing skills are never added as claims.
    generated_payload, generation_changes = build_optimized_cv_data(
        cv_payload, payload.target_text, opt_result.matched_skills
    )

    # The optional Qwen model is a separate generation model from the BGE
    # matching model. Its output is constrained to existing entries and falls
    # back to the deterministic result when unavailable or invalid.
    llm_payload = llm_service.generate_optimized_cv(generated_payload, payload.target_text)
    if llm_payload:
        generated_payload = llm_payload
        generation_source = "qwen-generation"
    else:
        generation_source = "rule-based"
    generated_payload["contact"] = contact
    cv_data = parse_cv_data(generated_payload)

    try:
        safe_name = (contact.get("full_name") or "Optimized_CV").replace(" ", "_")
        if payload.format == "docx":
            doc_bytes = generate_docx(cv_data, template=payload.template)
            return Response(
                content=doc_bytes,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}_CV.docx"'},
            )
        else:
            pdf_bytes = generate_pdf(cv_data, template=payload.template)
            return Response(
                content=pdf_bytes,
                media_type="application/pdf",
                headers={"Content-Disposition": f'attachment; filename="{safe_name}_CV.pdf"'},
            )
    except Exception as e:
        logger.error(f"Optimized CV download failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")


# ---------------------------------------------------------------------------
# CV Generator API Routes (DOCX / PDF with proper formatting)
# ---------------------------------------------------------------------------


@app.get("/api/cv/generator/templates")
def list_generator_templates() -> dict:
    """List available professional CV templates for the generator."""
    return {"templates": AVAILABLE_TEMPLATES}


class CVGenerateRequest(BaseModel):
    """Request body for CV generation."""
    contact: dict = Field(default_factory=dict)
    summary: str = Field(default="")
    experience: list[dict] = Field(default_factory=list)
    education: list[dict] = Field(default_factory=list)
    skills: dict = Field(default_factory=dict)
    projects: list[dict] = Field(default_factory=list)
    certifications: list[dict] = Field(default_factory=list)
    languages: list[str] = Field(default_factory=list)
    interests: list[str] = Field(default_factory=list)
    custom_sections: dict = Field(default_factory=dict)
    template: str = Field(default="modern", max_length=30)


@app.post("/api/cv/generate/pdf")
def generate_cv_pdf(payload: CVGenerateRequest):
    """Generate a professionally formatted PDF CV.

    Accepts structured CV data and returns a downloadable PDF file with
    proper headings, bullet points, alignment, and styling.
    """
    try:
        cv_data = parse_cv_data(payload.model_dump())
        pdf_bytes = generate_pdf(cv_data, template=payload.template)
        filename = f"{cv_data.contact.full_name or 'CV'}_Resume.pdf".replace(" ", "_")
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"PDF generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@app.post("/api/cv/generate/docx")
def generate_cv_docx(payload: CVGenerateRequest):
    """Generate a professionally formatted DOCX CV.

    Accepts structured CV data and returns a downloadable Word document with
    proper headings, bullet points, alignment, and styling.
    """
    try:
        cv_data = parse_cv_data(payload.model_dump())
        docx_bytes = generate_docx(cv_data, template=payload.template)
        filename = f"{cv_data.contact.full_name or 'CV'}_Resume.docx".replace(" ", "_")
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


@app.post("/api/models/upload")
def upload_model(file: UploadFile = File(...), name: str | None = Form(None)) -> dict:
    """Upload a packaged model (zip or folder) into runtime/models.

    If a zip archive is uploaded it will be extracted to runtime/models/<name>.
    """
    MODEL_DIR.mkdir(parents=True, exist_ok=True)
    filename = (name or file.filename).strip()
    if not filename:
        raise HTTPException(status_code=400, detail="Provide a valid filename or name form field.")

    dest = MODEL_DIR / filename
    try:
        # save uploaded file to a temp path
        data = file.file.read()
        temp_path = MODEL_DIR / (filename + ".upload")
        with temp_path.open("wb") as fh:
            fh.write(data)

        # if zip, extract
        if zipfile.is_zipfile(temp_path):
            extract_dir = MODEL_DIR / (Path(filename).stem)
            extract_dir.mkdir(parents=True, exist_ok=True)
            with zipfile.ZipFile(temp_path, "r") as zf:
                zf.extractall(path=extract_dir)
            temp_path.unlink()
            return {"uploaded": True, "path": str(extract_dir)}

        # if tar (tar.gz) extract
        if tarfile.is_tarfile(temp_path):
            extract_dir = MODEL_DIR / (Path(filename).stem)
            extract_dir.mkdir(parents=True, exist_ok=True)
            with tarfile.open(temp_path, "r:*") as tf:
                tf.extractall(path=extract_dir)
            temp_path.unlink()
            return {"uploaded": True, "path": str(extract_dir)}

        # otherwise move to a folder named filename
        final_dir = MODEL_DIR / filename
        final_dir.mkdir(parents=True, exist_ok=True)
        with (final_dir / "uploaded.bin").open("wb") as fh:
            fh.write(data)
        temp_path.unlink()
        return {"uploaded": True, "path": str(final_dir)}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Failed to save model: {exc}")


@app.post("/api/embeddings/precompute", response_model=EmbeddingPrecomputeResponse)
def precompute_embedding_cache(payload: EmbeddingPrecomputeRequest) -> dict:
    cached = precompute_embeddings(get_engine().semantic_model, [text.strip() for text in payload.texts if text.strip()])
    if not cached:
        raise HTTPException(status_code=400, detail="Provide at least one non-empty text.")
    return {"cached": cached}


@app.get("/api/models")
def list_models() -> dict:
    """List packaged models under runtime/models and whether the engine has loaded one."""
    engine = get_engine()
    models = []
    if MODEL_DIR.exists():
        for p in sorted(MODEL_DIR.iterdir()):
            if not p.exists():
                continue
            files = []
            if p.is_dir():
                for f in sorted(p.iterdir()):
                    files.append(f.name)
            else:
                files = [p.name]
            loaded = False
            # engine reports loaded packaged model as 'packaged:<name>'
            try:
                if engine.semantic_enabled and isinstance(engine.semantic_model_name, str) and engine.semantic_model_name.startswith("packaged:"):
                    loaded = engine.semantic_model_name.split(":", 1)[1] == p.name
            except Exception:
                loaded = False
            models.append({"name": p.name, "files": files, "loaded": loaded})
    return {"models": models}


@app.post("/api/analyze")
async def analyze(
    resumes: Annotated[list[UploadFile], File(description="Resume files in PDF, DOCX, TXT, or MD format.")],
    job_file: Annotated[UploadFile | None, File(description="Optional job description file.")] = None,
    job_text: Annotated[str | None, Form(description="Optional pasted job description text.")] = None,
    role: Annotated[str, Form(description="recruiter or job-seeker")] = "recruiter",
    tfidf_weight: Annotated[float, Form(description="Weight for TF-IDF score between 0 and 1.")] = 0.65,
) -> dict:
    if not resumes:
        raise HTTPException(status_code=400, detail="Upload at least one resume.")
    if job_file is None and not (job_text and job_text.strip()):
        raise HTTPException(status_code=400, detail="Provide a job description file or pasted job text.")
    if not 0 <= tfidf_weight <= 1:
        raise HTTPException(status_code=400, detail="tfidf_weight must be between 0 and 1.")
    if role not in {"recruiter", "job-seeker"}:
        raise HTTPException(status_code=400, detail="role must be recruiter or job-seeker.")

    if job_file is not None:
        job_document = await extract_upload(job_file)
    else:
        job_document = ExtractedDocument(filename="pasted-job-description.txt", text=job_text or "", file_type="txt")

    resume_documents = [await extract_upload(file) for file in resumes]
    empty = [document.filename for document in resume_documents if not document.text.strip()]
    if empty:
        raise HTTPException(status_code=400, detail=f"No extractable text found in: {', '.join(empty)}")

    session = get_engine().analyze(
        job=job_document,
        resumes=resume_documents,
        role=role,
        tfidf_weight=tfidf_weight,
    )
    STORE.save(session)
    return to_plain(session)


@app.get("/api/sessions")
def list_sessions(limit: int = 20) -> dict:
    """List past screening sessions."""
    bounded = min(max(limit, 1), 100)
    sessions = STORE.list_sessions(limit=bounded)
    return {"sessions": sessions}


@app.get("/api/sessions/{session_id}")
def get_session(session_id: str) -> dict:
    session = STORE.load(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found.")
    return session


@app.get("/api/sessions/{session_id}/export.csv", response_class=PlainTextResponse)
def export_csv(session_id: str) -> PlainTextResponse:
    session = STORE.load(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found.")
    return PlainTextResponse(
        session_to_csv(session),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="screening-{session_id}.csv"'},
    )


@app.post("/api/extract-text")
async def extract_text(file: UploadFile = File(...)) -> PlainTextResponse:
    """Extract text from an uploaded file (PDF, DOCX, TXT, MD).

    Uses the same extraction pipeline as the analyze endpoint.
    Returns plain text suitable for CV/JD input.
    """
    doc = await extract_upload(file)
    if not doc.text.strip():
        raise HTTPException(status_code=400, detail="No extractable text found in the uploaded file.")
    return PlainTextResponse(
        doc.text,
        headers={"Content-Type": "text/plain; charset=utf-8"},
    )


@app.post("/api/convert")
async def convert_to_pdf(file: UploadFile = File(...)) -> Response:
    """Convert an uploaded document (DOCX, TXT, MD) to PDF for preview.

    Uses python-docx to read .docx files and fpdf2 to generate a PDF.
    Falls back to plain-text PDF for unsupported formats.
    """
    from io import BytesIO

    raw = await file.read()
    filename = file.filename or "document"
    suffix = Path(filename).suffix.lower()

    text = ""
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
        except Exception:
            text = raw.decode("utf-8", errors="ignore")
    elif suffix in (".txt", ".md", ".csv"):
        text = raw.decode("utf-8", errors="ignore")
    elif suffix == ".pdf":
        # Already PDF — return as-is
        return Response(raw, media_type="application/pdf", headers={"Content-Disposition": f'inline; filename="{filename}"'})
    else:
        text = raw.decode("utf-8", errors="ignore")

    # Generate PDF from text using fpdf2
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", uni=True)
        pdf.set_font("DejaVu", "", 10)
        for line in text.split("\n"):
            line = line.strip()
            if line:
                # Handle non-Latin characters safely
                safe_line = line.encode("latin-1", "replace").decode("latin-1")
                pdf.cell(0, 5, safe_line, new_x="LMARGIN", new_y="NEXT")
            else:
                pdf.ln(3)

        pdf_bytes = pdf.output(dest="S").encode("latin-1")
        return Response(
            pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'inline; filename="{Path(filename).stem}.pdf"'},
        )
    except Exception as exc:
        # If PDF generation fails, return plain text
        return Response(
            text,
            media_type="text/plain",
            headers={"Content-Disposition": f'inline; filename="{Path(filename).stem}.txt"'},
        )


@app.post("/api/compare")
async def compare_resumes_api(
    resume_a: UploadFile = File(...),
    resume_b: UploadFile = File(...),
) -> dict:
    """Compare two resumes side-by-side: skill overlap, similarity score,
    tie-breaking recommendation, and duplicate detection.
    """
    doc_a = await extract_upload(resume_a)
    doc_b = await extract_upload(resume_b)

    result = compare_resumes(doc_a, doc_b)
    return result


@app.get("/api/sessions/{session_id}/candidates/{candidate_id}/report.pdf")
def transparency_report(session_id: str, candidate_id: str) -> Response:
    session = STORE.load(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found.")
    try:
        body = candidate_report_pdf(session, candidate_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="Candidate not found.") from None

    content_type = "application/pdf" if body.startswith(b"%PDF") else "text/plain"
    suffix = "pdf" if content_type == "application/pdf" else "txt"
    return Response(
        body,
        media_type=content_type,
        headers={"Content-Disposition": f'attachment; filename="{candidate_id}-transparency-report.{suffix}"'},
    )


@app.get("/api/sessions/{session_id}/candidates/{candidate_id}/cv.pdf")
def candidate_cv_pdf(session_id: str, candidate_id: str) -> Response:
    """Render the candidate's (anonymized) CV as an inline PDF for in-app viewing."""
    session = STORE.load(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found.")
    try:
        body = resume_to_pdf(session, candidate_id)
    except KeyError:
        raise HTTPException(status_code=404, detail="Candidate not found.") from None

    is_pdf = body.startswith(b"%PDF")
    return Response(
        body,
        media_type="application/pdf" if is_pdf else "text/plain; charset=utf-8",
        headers={"Content-Disposition": f'inline; filename="{candidate_id}-cv.{"pdf" if is_pdf else "txt"}"'},
    )


# ---------------------------------------------------------------------------
# Auth Routes — email/password + Google SSO (web-server code flow)
# ---------------------------------------------------------------------------


class SignupRequest(BaseModel):
    email: str = Field(min_length=3, max_length=120)
    password: str = Field(min_length=6, max_length=200)
    name: str = Field(default="", max_length=120)
    role: str = Field(default="candidate")


class LoginRequest(BaseModel):
    email: str = Field(min_length=3, max_length=120)
    password: str = Field(min_length=1, max_length=200)


class VerifyEmailRequest(BaseModel):
    email: str = Field(min_length=3, max_length=120)
    token: str | None = Field(default=None)
    code: str | None = Field(default=None)


class ResendRequest(BaseModel):
    email: str = Field(min_length=3, max_length=120)


async def _send_verification(email: str, name: str) -> None:
    """Generate a fresh verification code/token and email it via Resend."""
    pair = generate_verification(email)
    if pair is None:
        return
    code, token = pair
    try:
        await send_verification_email(email, name, code, token)
    except Exception as exc:  # never block signup on email failure
        logger.warning("Verification email failed for %s: %s", email, exc)


@app.post("/api/auth/signup")
async def signup(payload: SignupRequest) -> dict:
    """Create an account with email + password and send a verification email."""
    role = payload.role if payload.role in SELF_SERVICE_ROLES else "candidate"
    email = payload.email.lower().strip()
    if "@" not in email or "." not in email.split("@")[-1]:
        raise HTTPException(status_code=400, detail="Please enter a valid email address")
    try:
        user = create_user(
            email=email,
            password=payload.password,
            name=payload.name,
            role=role,
            email_verified=False,
        )
    except ValueError:
        raise HTTPException(status_code=409, detail="An account with this email already exists")

    await _send_verification(email, user["name"])

    token = authenticate(email, payload.password)
    return {
        **user,
        "token": token["token"] if token else None,
        "provider": "email",
    }


@app.post("/api/auth/login")
def login(payload: LoginRequest) -> dict:
    """Authenticate with email + password."""
    result = authenticate(payload.email.lower().strip(), payload.password)
    if result is None:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    result["provider"] = "email"
    return result


@app.get("/api/v1/auth/google/login")
def google_login(role: str = "candidate", mode: str = "login") -> RedirectResponse:
    """Redirect the browser to Google's OAuth consent screen."""
    if not google_configured():
        raise HTTPException(status_code=503, detail="Google sign-in is not configured")
    url = build_google_auth_url(role=role, mode=mode)
    return RedirectResponse(url, status_code=307)


@app.get("/api/v1/auth/google/callback")
async def google_callback(
    code: str | None = None,
    state: str = "",
    error: str | None = None,
) -> RedirectResponse:
    """Handle Google's redirect: exchange code, sign the user in, bounce to the frontend."""
    if error or not code:
        return RedirectResponse(f"{APP_URL}/login?error=google_denied", status_code=303)

    result = await google_oauth_callback(code=code, state=state)
    if result is None:
        return RedirectResponse(f"{APP_URL}/login?error=google_failed", status_code=303)

    # Pass the JWT back to the frontend callback page via the URL fragment-style query.
    params = urlencode({
        "token": result["token"],
        "role": result["role"],
        "new": "1" if result.get("is_new") else "0",
    })
    return RedirectResponse(f"{APP_URL}/callback?{params}", status_code=303)


@app.post("/api/auth/verify-email")
def verify_email(payload: VerifyEmailRequest) -> dict:
    """Verify a user's email via the link token or the 6-digit code."""
    if not payload.token and not payload.code:
        raise HTTPException(status_code=400, detail="A verification token or code is required")
    user = verify_email_record(payload.email.lower().strip(), token=payload.token, code=payload.code)
    if user is None:
        raise HTTPException(status_code=400, detail="Invalid or expired verification code")
    return {"verified": True, **user}


@app.post("/api/auth/resend-verification")
async def resend_verification(payload: ResendRequest) -> dict:
    """Re-send the verification email for an unverified account."""
    email = payload.email.lower().strip()
    user = get_user_by_email(email)
    if user is None:
        # Do not reveal whether the email exists.
        return {"sent": True}
    if user.get("email_verified"):
        return {"sent": True, "already_verified": True}
    await _send_verification(email, user.get("name", ""))
    return {"sent": True}


@app.get("/api/auth/me")
def me(current_user: dict = Depends(get_current_user)) -> dict:
    return current_user


@app.post("/api/auth/logout")
def logout() -> dict:
    # Stateless JWT — client discards the token
    return {"message": "Logged out successfully"}


@app.put("/api/auth/profile")
def update_profile(
    payload: dict,
    current_user: dict = Depends(get_current_user),
) -> dict:
    allowed = {"name", "password"}
    updates = {k: v for k, v in payload.items() if k in allowed and v}
    if not updates:
        raise HTTPException(status_code=400, detail="No updatable fields provided")
    user = update_user(current_user["user_id"], updates)
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")
    return user


# ---------------------------------------------------------------------------
# Job Board API Routes
# ---------------------------------------------------------------------------



class RoleChangeRequest(BaseModel):
    role: str = Field(pattern="^(candidate|recruiter)$")


@app.put("/api/auth/role")
def change_role(
    payload: RoleChangeRequest,
    current_user: dict = Depends(get_current_user),
) -> dict:
    """Persist a self-service role chosen after sign-in and refresh the token."""
    user = update_user(current_user["user_id"], {"role": payload.role})
    if user is None:
        raise HTTPException(status_code=404, detail="User not found")
    return {
        **user,
        "token": issue_token(user["user_id"], payload.role),
    }

@app.get("/api/external-jobs")
def list_external_jobs(
    limit: int = 50,
    source: str | None = None,
    query: str = "",
    current_user: dict = Depends(get_current_user),
) -> dict:
    """List imported external jobs with optional source and text search filters."""
    jobs = JOB_BOARD.list_jobs(limit=limit, source=source, query=query)
    return {"jobs": jobs, "total": len(jobs)}


@app.get("/api/external-jobs/sources")
def list_job_sources() -> dict:
    """List available job sources with configuration status."""
    return {
        "sources": [
            {
                "id": "greenhouse",
                "name": "Greenhouse",
                "configured": True,  # Public API, no key needed
                "docs_url": "https://developers.greenhouse.io/job-board.html",
                "needs_api_key": False,
            },
        ]
    }


class AdzunaImportRequest(BaseModel):
    query: str = Field(default="software engineer", max_length=200)
    location: str = Field(default="", max_length=100)
    country: str = Field(default="us", max_length=5)
    max_results: int = Field(default=50, ge=1, le=200)


@app.post("/api/external-jobs/import/adzuna")
async def import_adzuna_jobs(
    payload: AdzunaImportRequest,
    current_user: dict = Depends(get_current_user),
) -> dict:
    """Import jobs from Adzuna API."""
    if current_user.get("role") not in ("admin", "recruiter"):
        raise HTTPException(status_code=403, detail="Only recruiters and admins can import jobs")

    jobs = await adzuna_import(
        query=payload.query,
        location=payload.location,
        country=payload.country,
        max_results=payload.max_results,
    )

    imported = []
    for job in jobs:
        try:
            JOB_BOARD.save(job)
            imported.append(job.source_id)
        except Exception as exc:
            logger.warning("Failed to save Adzuna job: %s", exc)

    return {"imported": len(imported), "source": "adzuna", "job_ids": imported}


class GreenhouseImportRequest(BaseModel):
    board_tokens: list[str] = Field(min_length=1, max_length=50)
    max_per_company: int = Field(default=50, ge=1, le=100)


@app.post("/api/external-jobs/import/greenhouse")
async def import_greenhouse_jobs(
    payload: GreenhouseImportRequest,
    current_user: dict = Depends(get_current_user),
) -> dict:
    """Import jobs from Greenhouse company boards."""
    if current_user.get("role") not in ("admin", "recruiter"):
        raise HTTPException(status_code=403, detail="Only recruiters and admins can import jobs")

    jobs = await greenhouse_import(
        board_tokens=payload.board_tokens,
        max_per_company=payload.max_per_company,
    )

    imported = []
    for job in jobs:
        try:
            JOB_BOARD.save(job)
            imported.append(job.source_id)
        except Exception as exc:
            logger.warning("Failed to save Greenhouse job: %s", exc)

    return {"imported": len(imported), "source": "greenhouse", "job_ids": imported}


@app.delete("/api/external-jobs/{job_id}")
def delete_external_job(
    job_id: str,
    current_user: dict = Depends(get_current_user),
) -> dict:
    """Delete an imported external job."""
    if current_user.get("role") not in ("admin", "recruiter"):
        raise HTTPException(status_code=403, detail="Only recruiters and admins can delete jobs")
    job = JOB_BOARD.load(job_id)
    if job is None:
        raise HTTPException(status_code=404, detail="Job not found")
    JOB_BOARD.delete(job_id)
    return {"deleted": True, "job_id": job_id}